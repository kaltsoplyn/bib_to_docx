import sys
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import RGBColor
from bib_to_dict import BibToDictList, BibRecord

# This script parses a bibliography export from WoK, into a docx file with list of bib entries.

if (len(sys.argv) > 1):
    filename = sys.argv[1]
else: 
    filename = 'bib_in.bib'

print("Parsing file: {}\n".format(filename))

bib_in = BibToDictList(filename)
records = bib_in.records

# function to gather info from one BibRecord -> returns a dict
def build_rec(rec):
    out = {
        "title": rec.get_tag('title'),
        "authors": rec.get_tag('author'),
        "journal_title": rec.get_tag('journal') + " ",
        "journal_vol": rec.get_tag('volume'),
        "journal_rest": (" (" + rec.get_tag('number') + ")" if rec.get_tag('number') != '' else '' ) \
                    + (", " + rec.get_tag('pages') ) \
                    + " (" + rec.get_tag('year') + ")",
        "doi": rec.get_tag('doi'),
        "doi_url": "https://dx.doi.org/{}".format(rec.get_tag('doi')),
        "oa": rec.get_tag('oa')
    }
    return out

# gather all record dicts in a list
all_recs = [build_rec(rec) for rec in records]


# Now, onto the docx file containing the list of publications

document = Document()
document.add_heading('Bibliography', 0)

Colors = {
    "Gold": RGBColor(0xd4, 0xaf, 0x37),
    "Green": RGBColor(0x22, 0xac, 0x00),
    "Black": RGBColor(0x00, 0x00, 0x00),
    "Gray": RGBColor(0xaa, 0xaa, 0xaa),
    "Blue": RGBColor(0x00, 0x99, 0xcc)
    }

# function to define styles
# USAGE: new_style( style name: string, font_size: int pt, kw:is_bold: bool, kwarg: is_italic: bool, kw: left_ind: float cm, kw: first_line_ind: float cm )
# pretty self explanatory except: left_ind is left_indent = how much indent from the left in cm
# the first_line_ind defines how much the first line is indented *RELATIVE* to the rest (aka relative to left_ind)
def new_style( name, font_size, font_name = 'Calibri', is_bold = False, is_italic = False, left_ind = 0.5, first_line_ind = 0, color = Colors["Black"]  ):
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = font_name
    style.font.bold = is_bold
    style.font.italic = is_italic
    style.font.size = Pt(font_size)
    # style.font.color.theme_color = color
    style.font.color.rgb = color
    paragraph_format = style.paragraph_format
    paragraph_format.first_line_indent = Cm(first_line_ind)
    paragraph_format.left_indent = Cm(left_ind)
    paragraph_format.space_before = Pt(4)
    paragraph_format.space_after = Pt(2)
    return style

# define styles
# the "DD" styles are defined with larger indents so that things align nicely when the index is > 10
# I duplicated the whole style in code, because my attempt to copy or deepcopy the styles and change just the indent failed for some reason
title_style = new_style('title_style', 11, is_bold=True, first_line_ind=-0.4 )
title_styleDD = new_style('title_styleDD', 11, is_bold=True, left_ind=0.5, first_line_ind=-0.6 )
authors_style = new_style('authors_style', 11, is_italic=True)
authors_styleDD = new_style('authors_styleDD', 11, is_italic=True, left_ind=0.5)
journal_style = new_style('journal_style', 11)
journal_styleDD = new_style('journal_styleDD', 11, left_ind=0.5)


# copy-pasted function to add hyperlinks, as, strangely, there is no direct way through the API
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    #r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.color.rgb = Colors["Blue"]
    r.font.underline = True

    return hyperlink

# define the function writing an entry to the docx file for each bib record
def build_entry(index, r):

    title = document.add_paragraph()
    title.style = title_style if index + 1 < 10 else title_styleDD

    title.add_run(str(index + 1) + ". " + r['title'] + "  ")
    
    if (r['oa'] == ''):
        OAdecorator = '🔒'
        OAcolor = Colors["Gray"]
    else:
        OAdecorator = '🔓'
        OAcolor = Colors["Gold"] if r['oa'] == 'gold' else Colors["Green"]

    OArun = title.add_run(OAdecorator)
    OArun.font.color.rgb = OAcolor
    OArun.font.name = 'Segoe UI Symbol'

    authors = document.add_paragraph()
    authors.style = authors_style if index + 1 < 10 else authors_styleDD
    authors.add_run(r['authors'].replace(', and ', ', ', r['authors'].count(', and ') - 1)).italic = True

    journal = document.add_paragraph()
    journal.style = journal_style if index + 1 < 10 else journal_styleDD
    journal.add_run(r['journal_title'])
    journal.add_run(r['journal_vol']).bold = True
    journal.add_run(r['journal_rest'] + "  |  doi: ")
    #journal.add_run(r['doi'])
    add_hyperlink(journal, r['doi'], r['doi_url'])

# run the build function over all records
for (index, rec) in enumerate(all_recs):
    build_entry(index, rec)

# save the docx
try:
    document.save('bib_out.docx')
    print("Presumably things worked as planned. Output file: 'bib_out.docx'")
except Exception as e:
    print("Exception occured with output file: {}".format(e))